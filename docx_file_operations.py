from docx import Document
import re

# Создаем новый документ
doc = Document()

# Добавляем текст в документ
doc.add_paragraph("Это содержимое файла B.")

# Сохраняем документ в файл формата DOCX
doc.save("B.docx")

print("Файл B.docx успешно создан.")

from docx import Document

# Открываем существующий документ формата DOCX
doc = Document("B.docx")

# Открываем исходный файл для чтения
with open('исходный_файл.txt', 'r') as file:
    text = file.read()

# Добавляем текст из исходного файла в документ
doc.add_paragraph(text)

# Сохраняем изменения в файле B.docx
doc.save("B.docx")

print("Файл B.docx успешно обновлен.")

# Выполнить поиск имен с большой буквы в файле B.docx
doc = Document("B.docx")
found_names = []

for paragraph in doc.paragraphs:
    matches = re.findall(r'\b[A-Z][a-zA-Z]+\b', paragraph.text)
    found_names.extend(matches)

if found_names:
    print("Найденные имена с большой буквы:")
    for name in found_names:
        print(name)
else:
    print("В файле B.docx не найдены имена с большой буквы.")

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK

# Создаем новый документ
doc = Document()

# Устанавливаем нестандартный стиль
custom_style = doc.styles.add_style('CustomStyle', WD_PARAGRAPH_ALIGNMENT.CENTER)
custom_style.font.size = Pt(16)
custom_style.font.bold = True

# Добавляем заголовок
doc.add_paragraph("Заголовок документа", style='CustomStyle')

# Добавляем разрыв страницы
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# Добавляем текст с разрывом строк
doc.add_paragraph("Первый абзац.")
doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
doc.add_paragraph("Второй абзац.")

# Добавляем разрыв страницы
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# Добавляем изображение
doc.add_picture('image.jpg')

# Сохраняем документ с нестандартным именем файла
file_name = "CustomDocument.docx"
doc.save(file_name)

print(f"Документ {file_name} успешно создан.")

input()
